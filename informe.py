import tkinter as tk
from tkinter import ttk, messagebox
import os
import webbrowser
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==============================================================================================
# 1. EL MOTOR DEL DOCUMENTO (Blindado con la rúbrica y Matriz ESA Unificada)
# ==============================================================================================
class GeneradorInforme:
    def __init__(self, titulo_proyecto, numero_grupo):
        self.doc = Document()
        self.titulo = titulo_proyecto
        self.grupo = numero_grupo
        self.configurar_formato_rubrica()
        self.configurar_paginacion()

    def configurar_paginacion(self):
        section = self.doc.sections[0]
        section.footer_distance = Cm(1.5)
        footer = section.footer
        p_footer = footer.paragraphs[0]
        p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = p_footer.add_run()
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.extend([fldChar1, instrText, fldChar2, fldChar3])
        run.font.name = 'Arial'
        run.font.size = Pt(12)

    def configurar_formato_rubrica(self):
        section = self.doc.sections[0]
        section.page_height = Cm(27.94) 
        section.page_width = Cm(21.59)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        style_normal = self.doc.styles['Normal']
        style_normal.font.name = 'Arial'
        style_normal.font.size = Pt(12)
        style_normal.paragraph_format.line_spacing = 1.5
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        h1 = self.doc.styles['Heading 1']
        h1.font.name, h1.font.size, h1.font.bold, h1.font.color.rgb = 'Arial', Pt(12), True, RGBColor(0,0,0)
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h1.paragraph_format.page_break_before = True 

        h2 = self.doc.styles['Heading 2']
        h2.font.name, h2.font.size, h2.font.bold, h2.font.color.rgb = 'Arial', Pt(12), True, RGBColor(0,0,0)
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2.paragraph_format.left_indent = Cm(1.25)

        for name in ['LeyendaFigura', 'LeyendaTabla']:
            if name not in self.doc.styles:
                style = self.doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
                style.font.name = 'Arial'
                style.font.size = Pt(9)
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                style.paragraph_format.line_spacing = 1.0 

        try:
            toc1 = self.doc.styles['TOC 1']
        except KeyError:
            toc1 = self.doc.styles.add_style('TOC 1', WD_STYLE_TYPE.PARAGRAPH)
        
        toc1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        toc1.paragraph_format.left_indent = Cm(0) 
        toc1.font.name = 'Arial'
        toc1.font.size = Pt(12)
        toc1.font.color.rgb = RGBColor(0,0,0)

    def _insertar_campo_indice(self, instruccion_xml):
        parrafo = self.doc.add_paragraph()
        run = parrafo.add_run()
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
        instrText.text = instruccion_xml
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

    def _set_cell_background(self, celda, hex_color):
        tcPr = celda._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def generar_portada(self, ruta_logo, facultad, curso, autores, profesor, ciudad, anio):
        extensiones = ['', '.png', '.jpg', '.jpeg', '.gif', '.bmp']
        ruta_final_logo = None
        for ext in extensiones:
            if os.path.exists(ruta_logo + ext):
                ruta_final_logo = ruta_logo + ext
                break

        p_logo = self.doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ruta_final_logo:
            p_logo.add_run().add_picture(ruta_final_logo, width=Cm(4.5))

        p_inst = self.doc.add_paragraph(); p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_inst.add_run(facultad)
        r.add_break()
        r.add_text(curso)

        self.doc.add_paragraph()
        p_tit = self.doc.add_paragraph(); p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_tit.add_run(self.titulo).bold = True

        self.doc.add_paragraph()
        p_aut = self.doc.add_paragraph(); p_aut.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_aut.add_run("Autores:")
        r.add_break()
        autores_limpios = [a.strip() for a in autores if a and a.strip()]
        if autores_limpios:
            def _clave_orden(a):
                partes = a.split()
                return partes[-2].lower() if len(partes) >= 2 else a.lower()
            for autor in sorted(autores_limpios, key=_clave_orden):
                r = p_aut.add_run(autor)
                r.add_break()

        self.doc.add_paragraph()
        p_prof = self.doc.add_paragraph(); p_prof.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_prof.add_run("Profesor:")
        r.add_break()
        r.add_text(profesor)

        self.doc.add_paragraph()
        p_pie = self.doc.add_paragraph(); p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_pie.add_run(f"{ciudad}, Chile")
        r.add_break()
        r.add_text(anio)

    def estructurar_indices(self):
        for titulo, xml in [
            ('ÍNDICE DE CONTENIDOS', 'TOC \\o "1-3" \\h \\z \\u'),
            ('ÍNDICE DE FIGURAS', 'TOC \\h \\z \\t "LeyendaFigura;1"'),
            ('ÍNDICE DE TABLAS', 'TOC \\h \\z \\t "LeyendaTabla;1"')
        ]:
            p = self.doc.add_heading(titulo, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._insertar_campo_indice(xml)

    def agregar_capitulo(self, titulo, texto=None):
        p_t = self.doc.add_heading(titulo.upper(), level=1)
        self.doc.add_paragraph() 
        self.doc.add_paragraph() 
        if texto: self.doc.add_paragraph(texto)

    def _formatear_subtitulo(self, texto):
        if not texto: return ""
        for i, char in enumerate(texto):
            if char.isalpha():
                return texto[:i] + char.upper() + texto[i+1:].lower()
        return texto

    def agregar_subtitulo(self, subtitulo, texto=None):
        sub_formateado = self._formatear_subtitulo(subtitulo)
        p_t = self.doc.add_heading(sub_formateado, level=2)
        self.doc.add_paragraph() 
        if texto: self.doc.add_paragraph(texto)

    def agregar_texto_libre(self, texto, negrita=False, cursiva=False):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(texto)
        run.bold, run.italic = negrita, cursiva
        self.doc.add_paragraph() 

    def agregar_definicion(self, termino, descripcion):
        p_termino = self.doc.add_paragraph()
        p_termino.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_termino.add_run(termino).bold = True
        p_desc = self.doc.add_paragraph(descripcion)
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph()

    def agregar_bibliografia_apa(self, lista_referencias):
        self.doc.add_heading("REFERENCIAS BIBLIOGRÁFICAS", level=1)
        self.doc.add_paragraph()
        for ref in lista_referencias:
            p = self.doc.add_paragraph(ref)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-1.25) 
            p.paragraph_format.line_spacing = 1.0 
            p.paragraph_format.space_after = Pt(12)

    def agregar_figura(self, ruta_imagen, correlativo, leyenda):
        extensiones = ['', '.png', '.jpg', '.jpeg', '.gif', '.bmp']
        ruta_final = None
        for ext in extensiones:
            if os.path.exists(ruta_imagen + ext):
                ruta_final = ruta_imagen + ext
                break
                
        p_tit = self.doc.add_paragraph(style='LeyendaFigura')
        p_tit.add_run(f"Figura {correlativo} ")
        p_tit.add_run(f'"{leyenda}".').italic = True
        
        p_img = self.doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ruta_final:
            p_img.add_run().add_picture(ruta_final, width=Cm(14))
        
        p_f = self.doc.add_paragraph("Fuente: Elaborado por el estudiante de acuerdo con el proyecto.")
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f.paragraph_format.line_spacing = 1.0
        p_f.runs[0].font.name, p_f.runs[0].font.size = 'Arial', Pt(9)

    def _crear_titulo_tabla(self, correlativo, leyenda):
        p_tit = self.doc.add_paragraph(style='LeyendaTabla')
        p_tit.add_run(f"Tabla {correlativo}: ")
        p_tit.add_run(f'"{leyenda}".').italic = True

    def _crear_fuente_tabla(self):
        p_f = self.doc.add_paragraph("Fuente: Elaborado por el estudiante de acuerdo con el proyecto.")
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f.paragraph_format.line_spacing = 1.0 
        p_f.runs[0].font.name, p_f.runs[0].font.size = 'Arial', Pt(9)
        self.doc.add_paragraph()

    def agregar_tabla(self, encabezados, filas, correlativo, leyenda, color_encabezado=None):
        self._crear_titulo_tabla(correlativo, leyenda)
        tabla = self.doc.add_table(rows=1, cols=len(encabezados)); tabla.style = 'Table Grid'
        
        for i, h in enumerate(encabezados):
            c = tabla.rows[0].cells[i]
            c.text = str(h)
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if color_encabezado: self._set_cell_background(c, color_encabezado) 
                
        for fila in filas:
            cells = tabla.add_row().cells
            for i, dato in enumerate(fila): 
                cells[i].text = str(dato)
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
        self._crear_fuente_tabla()

    def agregar_caso_uso_extendido(self, nombre, actores, resumen, frecuencia, precondiciones, descripcion, excepciones, poscondiciones, dependencias, correlativo):
        self._crear_titulo_tabla(correlativo, f"Caso de Uso: {nombre}")
        
        datos = [
            ("Nombre", nombre), ("Actor(es)", actores), ("Resumen", resumen),
            ("Frecuencia", frecuencia), ("Precondiciones", precondiciones), 
            ("Descripción", descripcion), ("Excepciones", excepciones), 
            ("Poscondiciones", poscondiciones), ("Dependencias", dependencias)
        ]
        
        tabla = self.doc.add_table(rows=1, cols=2)
        tabla.style = 'Table Grid'
        
        c0 = tabla.rows[0].cells[0]
        c0.text = "Campo"; c0.paragraphs[0].runs[0].bold = True; c0.width = Cm(4.5)
        self._set_cell_background(c0, "D9D9D9")
        
        c1 = tabla.rows[0].cells[1]
        c1.text = "Descripción"; c1.paragraphs[0].runs[0].bold = True
        self._set_cell_background(c1, "D9D9D9")
        
        for etiqueta, valor in datos:
            row = tabla.add_row()
            c_izq = row.cells[0]; c_izq.text = etiqueta; c_izq.paragraphs[0].runs[0].bold = True; c_izq.width = Cm(4.5)
            self._set_cell_background(c_izq, "F2F2F2") 
            
            c_der = row.cells[1]
            lineas = str(valor).split('\n')
            c_der.text = lineas[0]
            c_der.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for linea in lineas[1:]:
                p = c_der.add_paragraph(linea)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
        self._crear_fuente_tabla()

    def agregar_matriz_ur_esa(self, categorias, correlativo):
        self._crear_titulo_tabla(correlativo, "Requerimientos funcionales formato ESA")
        tabla = self.doc.add_table(rows=1, cols=9); tabla.style = 'Table Grid'
        
        headers = ["", "ID", "Descripción", "Necesidad", "Prioridad", "Estabilidad", "Claridad", "Verificabilidad", "Fuente"]
        for i, h in enumerate(headers):
            c = tabla.rows[0].cells[i]; c.text = h; c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for cat_idx, cat in enumerate(categorias):
            row_c = tabla.add_row()
            row_c.cells[0].text = f"{cat_idx + 1}."
            row_c.cells[0].paragraphs[0].runs[0].bold = True
            c_nom = row_c.cells[2]; c_nom.merge(row_c.cells[-1]) 
            c_nom.text = cat.get("nombre", ""); c_nom.paragraphs[0].runs[0].bold = True
            
            for req in cat.get("reqs", []):
                r_d = tabla.add_row()
                r_d.cells[0].text = req.get("tipo", "UR"); r_d.cells[0].paragraphs[0].runs[0].bold = True; r_d.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_d.cells[1].text = req.get("id", ""); r_d.cells[1].paragraphs[0].runs[0].bold = True; r_d.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_d.cells[2].text = req.get("nombre", ""); r_d.cells[2].paragraphs[0].runs[0].bold = True; r_d.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for idx_p, pto in enumerate(req.get("puntos", [])):
                    if 3+idx_p < 9: r_d.cells[3+idx_p].text = str(pto); r_d.cells[3+idx_p].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                r_des = tabla.add_row(); c_desc = r_des.cells[1]; c_desc.merge(r_des.cells[-1]); c_desc.text = req.get("descripcion", ""); c_desc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
        self._crear_fuente_tabla()

    def compilar(self):
        nombre = f"Grupo {self.grupo} - Documento_Final.docx"
        self.doc.save(nombre)
        return nombre

# ==============================================================================================
# 2. LA INTERFAZ GRÁFICA (MATERIAL DESIGN PURO + MOCKUPS)
# ==============================================================================================
class AppGeneradorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Ingecon Document Studio 2026")
        self.root.geometry("1200x800")
        
        self.BG_MAIN = "#F0F2F5"      
        self.BG_CARD = "#FFFFFF"      
        self.C_PRIMARY = "#2563EB"    
        self.C_PRIMARY_H = "#1D4ED8"  
        self.C_SUCCESS = "#10B981"    
        self.C_SUCCESS_H = "#059669"
        self.C_WARNING = "#F59E0B"    
        self.C_WARNING_H = "#D97706"
        self.C_DANGER = "#EF4444"     
        self.C_DANGER_H = "#DC2626"
        self.C_TEXT = "#1F2937"       

        self.root.configure(bg=self.BG_MAIN)
        self.elementos_documento = []

        style = ttk.Style()
        style.theme_use('clam')
        style.configure(".", background=self.BG_MAIN, foreground=self.C_TEXT, font=("Segoe UI", 10))
        style.configure("TFrame", background=self.BG_CARD)
        style.configure("Card.TFrame", background=self.BG_CARD)
        style.configure("TNotebook", background=self.BG_MAIN, borderwidth=0)
        style.configure("TNotebook.Tab", background="#E5E7EB", foreground="#6B7280", padding=[20, 10], font=("Segoe UI", 11, "bold"), borderwidth=0)
        style.map("TNotebook.Tab", background=[("selected", self.BG_CARD)], foreground=[("selected", self.C_PRIMARY)])
        style.configure("TLabel", background=self.BG_CARD, foreground=self.C_TEXT, font=("Segoe UI", 10))
        style.configure("TEntry", padding=6, relief="flat", fieldbackground="#F9FAFB")

        self.notebook = ttk.Notebook(root)
        self.notebook.pack(expand=True, fill='both', padx=20, pady=20)
        
        self.tab_portada = ttk.Frame(self.notebook, style="Card.TFrame")
        self.tab_constructor = ttk.Frame(self.notebook, style="Card.TFrame")
        self.tab_preview = ttk.Frame(self.notebook, style="Card.TFrame")
        
        self.notebook.add(self.tab_portada, text=' 1. Datos del Proyecto ')
        self.notebook.add(self.tab_constructor, text=' 2. Constructor Visual ')
        self.notebook.add(self.tab_preview, text=' 3. Vista Previa A4 ')
        
        self.notebook.bind("<<NotebookTabChanged>>", self._al_cambiar_pestana)

        self._construir_tab_portada()
        self._construir_tab_constructor()
        self._construir_motor_grafico()
        self._cargar_plantilla_base()

    def _btn(self, parent, text, cmd, bg, hover_bg, fg="white", font=("Segoe UI", 10, "bold"), pady=10):
        b = tk.Button(parent, text=text, command=cmd, bg=bg, fg=fg, font=font, relief="flat", borderwidth=0, cursor="hand2", pady=pady)
        b.bind("<Enter>", lambda e: b.config(bg=hover_bg))
        b.bind("<Leave>", lambda e: b.config(bg=bg))
        return b

    def _cargar_plantilla_base(self):
        self._guardar_bloque(None, "📘 CAPÍTULO: Análisis Interno Externo", "capitulo", {"titulo": "Análisis Interno Externo de la Organización"})
        self._guardar_bloque(None, "   ↳ Subtítulo: Análisis PESTEL", "subtitulo", {"subtitulo": "Análisis PESTEL"})
        self._guardar_bloque(None, "📝 Texto: Para llevar a cabo el análisis...", "texto", {"texto": "Para llevar a cabo el análisis externo del macroentorno se aplicó la herramienta PESTEL de forma rigurosa..."})
        self._guardar_bloque(None, "📘 CAPÍTULO: Requerimientos del Sistema", "capitulo", {"titulo": "Requerimientos del Sistema"})
        self._guardar_bloque(None, "⚙️ Caso de Uso: CU-01 Iniciar Sesión", "caso_uso", {
            "nombre": "CU-01: Iniciar Sesión", "actores": "Administrador", "resumen": "Permite acceder a la plataforma.",
            "frecuencia": "Alta", "precondiciones": "Usuario registrado", "descripcion": "1. Ingresar credenciales.\n2. Validar en BD.",
            "excepciones": "Error 401", "poscondiciones": "Acceso al panel", "dependencias": "Ninguna"
        })
        self._guardar_bloque(None, "📋 MATRIZ UR ESA (1 Categorías)", "matriz_ur_esa", {
            "categorias": [{"nombre": "Formulario de Contacto: Visitante", "reqs": [
                {"tipo": "UR", "id": "1.1", "nombre": "Enviar Consulta", "puntos": ["1","1","1","1","1","1"], "descripcion": "El sistema debe permitir al Visitante enviar una Consulta de contacto..."}
            ]}]
        })

    def _crear_campo(self, parent, label_text, row, default_value=""):
        ttk.Label(parent, text=label_text, font=('Segoe UI', 10, 'bold')).grid(row=row, column=0, sticky='w', pady=8, padx=15)
        entry = ttk.Entry(parent, width=60)
        entry.insert(0, default_value)
        entry.grid(row=row, column=1, sticky='w', pady=8, padx=15)
        return entry

    def _construir_tab_portada(self):
        header = tk.Label(self.tab_portada, text="Configuración de la Portada", font=("Segoe UI", 16, "bold"), bg=self.BG_CARD, fg=self.C_PRIMARY)
        header.pack(anchor='w', padx=20, pady=20)
        frame = ttk.Frame(self.tab_portada)
        frame.pack(fill='both', expand=True, padx=20)
        
        self.ent_titulo = self._crear_campo(frame, "Título del Proyecto:", 0, "PROYECTO DE SOFTWARE INGECON")
        self.ent_grupo = self._crear_campo(frame, "N° de Grupo:", 1, "1")
        self.ent_logo = self._crear_campo(frame, "Nombre Logo (ej: logo_unab):", 2, "logo_unab")
        self.ent_facultad = self._crear_campo(frame, "Facultad:", 3, "Facultad de Ingeniería")
        self.ent_curso = self._crear_campo(frame, "Curso:", 4, "Ingeniería de Software I")
        self.ent_profesor = self._crear_campo(frame, "Profesor:", 5, "Paulo Quinsacara Jofré")
        self.ent_ciudad = self._crear_campo(frame, "Ciudad:", 6, "Santiago")
        self.ent_anio = self._crear_campo(frame, "Año:", 7, "2026")
        
        ttk.Label(frame, text="Autores (Uno por línea):", font=('Segoe UI', 10, 'bold')).grid(row=8, column=0, sticky='nw', pady=8, padx=15)
        self.txt_autores = tk.Text(frame, width=45, height=6, font=('Segoe UI', 10), relief="flat", highlightthickness=1, highlightbackground="#D1D5DB", highlightcolor=self.C_PRIMARY)
        self.txt_autores.insert('1.0', "Diego Carabajal Belmar\nTomás Hurtado Silva")
        self.txt_autores.grid(row=8, column=1, sticky='w', pady=8, padx=15)

    def _construir_tab_constructor(self):
        panel_izq = tk.Frame(self.tab_constructor, bg="#F9FAFB", width=260)
        panel_izq.pack(side=tk.LEFT, fill='y')
        tk.Label(panel_izq, text="Bloques", font=("Segoe UI", 12, "bold"), bg="#F9FAFB", fg=self.C_TEXT).pack(pady=(20, 10))

        # ESTRUCTURA DE BOTONES CON MOCKUPS AL LADO
        botones = [
            ("Añadir Capítulo", self._dlg_capitulo, "capitulo"), 
            ("Añadir Subtítulo", self._dlg_subtitulo, "subtitulo"),
            ("Añadir Texto Libre", self._dlg_texto, "texto"), 
            ("Añadir Figura", self._dlg_figura, "figura"),
            ("Añadir Tabla Simple", self._dlg_tabla_simple, "tabla"), 
            ("Añadir Caso Uso", self._dlg_caso_uso, "caso_uso"),
            ("Añadir Matriz UR ESA", self._dlg_matriz_ur, "matriz_ur_esa"), 
            ("Añadir Bibliografía", self._dlg_apa, "apa")
        ]
        
        for texto, comando, tipo in botones:
            f_row = tk.Frame(panel_izq, bg="#F9FAFB")
            f_row.pack(pady=4, padx=15, fill='x')
            
            b = self._btn(f_row, texto, comando, "#E5E7EB", "#D1D5DB", fg=self.C_TEXT, font=("Segoe UI", 10))
            b.pack(side=tk.LEFT, fill='x', expand=True)
            
            b_info = self._btn(f_row, "👁️", lambda t=tipo: self._mostrar_mockup(t), "#D1D5DB", "#9CA3AF", fg=self.C_TEXT, font=("Segoe UI", 10))
            b_info.pack(side=tk.RIGHT, padx=(5,0))

        panel_der = ttk.Frame(self.tab_constructor)
        panel_der.pack(side=tk.RIGHT, fill='both', expand=True, padx=20, pady=20)
        tk.Label(panel_der, text="Estructura del Documento", font=("Segoe UI", 16, "bold"), bg=self.BG_CARD, fg=self.C_PRIMARY).pack(anchor='w', pady=(0,10))

        frame_lista = tk.Frame(panel_der, bg=self.BG_CARD)
        frame_lista.pack(side=tk.LEFT, fill='both', expand=True)
        self.listbox = tk.Listbox(frame_lista, font=('Segoe UI', 11), relief="flat", highlightthickness=1, highlightbackground="#D1D5DB", highlightcolor=self.C_PRIMARY, selectbackground=self.C_PRIMARY, activestyle="none")
        self.listbox.pack(side=tk.LEFT, fill='both', expand=True)
        scrollbar = ttk.Scrollbar(frame_lista, orient="vertical", command=self.listbox.yview)
        scrollbar.pack(side=tk.LEFT, fill='y')
        self.listbox.config(yscrollcommand=scrollbar.set)
        
        # EVENTO DE DOBLE CLIC 
        self.listbox.bind("<Double-1>", lambda event: self._editar_elemento())

        frame_controles = tk.Frame(panel_der, bg=self.BG_CARD)
        frame_controles.pack(side=tk.RIGHT, fill='y', padx=(15,0))
        
        b_up = self._btn(frame_controles, "↑ Subir", self._mover_arriba, "#F3F4F6", "#E5E7EB", fg=self.C_TEXT)
        b_up.pack(pady=(0,5), fill='x')
        b_dn = self._btn(frame_controles, "↓ Bajar", self._mover_abajo, "#F3F4F6", "#E5E7EB", fg=self.C_TEXT)
        b_dn.pack(pady=5, fill='x')
        b_edit = self._btn(frame_controles, "✏️ Editar", self._editar_elemento, self.C_WARNING, self.C_WARNING_H)
        b_edit.pack(pady=15, fill='x')
        b_del = self._btn(frame_controles, "✖ Eliminar", self._eliminar_elemento, self.C_DANGER, self.C_DANGER_H)
        b_del.pack(pady=5, fill='x')
        b_gen = self._btn(frame_controles, "🚀 GENERAR", self.generar_documento, self.C_SUCCESS, self.C_SUCCESS_H, font=("Segoe UI", 12, "bold"), pady=15)
        b_gen.pack(side=tk.BOTTOM, pady=0, fill='x')

    # ================= FUNCIONES DE MOCKUPS VISUALES =================
    def _mostrar_mockup(self, tipo):
        win = tk.Toplevel(self.root)
        win.title("Vista Previa (Formato Rúbrica)")
        win.configure(bg="white")
        win.geometry("600x400")
        
        if tipo == "capitulo":
            tk.Label(win, text="2. CONTEXTO DE LA ORGANIZACIÓN", font=("Times New Roman", 16, "bold"), bg="white").pack(pady=(100,0))
        
        elif tipo == "subtitulo":
            tk.Label(win, text="2.1 Información general de la empresa", font=("Times New Roman", 14, "bold"), bg="white").pack(pady=(100,0), anchor="w", padx=50)
            
        elif tipo == "texto":
            texto = "El sistema debe permitir al Visitante enviar una Consulta de contacto, asegurando que la información sea registrada en la Base de Datos para su posterior gestión. Esto garantizará el flujo correcto de los requerimientos..."
            tk.Label(win, text=texto, font=("Arial", 11), bg="white", wraplength=500, justify="left").pack(pady=(100,0), padx=50)
            
        elif tipo == "figura":
            tk.Label(win, text="Figura 2.2 \"Organigrama Ingecon\".", font=("Times New Roman", 10, "italic"), bg="white").pack(pady=(50,10))
            f = tk.Frame(win, bg="#DBEAFE", width=400, height=200); f.pack()
            tk.Label(f, text="[ IMAGEN ]", bg="#DBEAFE", fg="#1E40AF", font=("Arial", 12, "bold")).place(relx=0.5, rely=0.5, anchor="center")
            tk.Label(win, text="Fuente: Elaborado por el estudiante de acuerdo con el proyecto.", font=("Times New Roman", 9), bg="white").pack(pady=(10,0))
            
        elif tipo == "tabla":
            tk.Label(win, text="Tabla 2.2: \"Requerimientos de admin\".", font=("Times New Roman", 10, "italic"), bg="white").pack(pady=(50,10))
            f = tk.Frame(win, bg="black", pady=1, padx=1); f.pack()
            tk.Label(f, text="ID", width=10, bg="#D9D9D9", font=("Arial", 10, "bold")).grid(row=0, column=0, padx=1, pady=1)
            tk.Label(f, text="Tarea", width=20, bg="#D9D9D9", font=("Arial", 10, "bold")).grid(row=0, column=1, padx=1, pady=1)
            tk.Label(f, text="1", width=10, bg="white", font=("Arial", 10)).grid(row=1, column=0, padx=1, pady=1)
            tk.Label(f, text="Diseño Base", width=20, bg="white", font=("Arial", 10)).grid(row=1, column=1, padx=1, pady=1)
            tk.Label(win, text="Fuente: Elaborado por el estudiante de acuerdo con el proyecto.", font=("Times New Roman", 9), bg="white").pack(pady=(10,0))
            
        elif tipo == "caso_uso":
            tk.Label(win, text="Tabla 9.5 CU 1.5 \"Seleccionando fecha de consulta\"", font=("Times New Roman", 10, "italic"), bg="white").pack(pady=(30,10))
            f = tk.Frame(win, bg="black", pady=1, padx=1); f.pack()
            tk.Label(f, text="Campo", width=15, bg="#D9D9D9", font=("Arial", 10, "bold")).grid(row=0, column=0, padx=1, pady=1, sticky="w")
            tk.Label(f, text="Descripción", width=40, bg="#D9D9D9", font=("Arial", 10, "bold")).grid(row=0, column=1, padx=1, pady=1, sticky="w")
            
            campos = ["Nombre", "Actor(es)", "Resumen", "Frecuencia", "Precondiciones"]
            for i, c in enumerate(campos):
                tk.Label(f, text=c, width=15, bg="#F2F2F2", font=("Arial", 10, "bold")).grid(row=i+1, column=0, padx=1, pady=1, sticky="w")
                tk.Label(f, text="...", width=40, bg="white", font=("Arial", 10)).grid(row=i+1, column=1, padx=1, pady=1, sticky="w")
            tk.Label(win, text="Fuente: Elaborado por el estudiante de acuerdo con el proyecto.", font=("Times New Roman", 9), bg="white").pack(pady=(10,0))

        elif tipo == "matriz_ur_esa":
            tk.Label(win, text="Tabla 8.8 \"Requerimientos funcionales formato ESA\".", font=("Times New Roman", 10, "italic"), bg="white").pack(pady=(30,10))
            f = tk.Frame(win, bg="black", pady=1, padx=1); f.pack(fill='x', padx=10)
            
            h = ["", "ID", "Descripción", "Nec", "Prio", "Est", "Cla", "Ver", "Fue"]
            for i, txt in enumerate(h):
                tk.Label(f, text=txt, width=15 if i==2 else 5, bg="white", font=("Times New Roman", 9, "bold")).grid(row=0, column=i, padx=1, pady=1)
                
            tk.Label(f, text="1.", bg="white", font=("Times New Roman", 9, "bold")).grid(row=1, column=0, padx=1, pady=1)
            tk.Label(f, text="Formulario de Contacto: Visitante", bg="white", font=("Times New Roman", 9, "bold"), anchor="w").grid(row=1, column=2, columnspan=7, sticky="we", padx=1, pady=1)
            
            r = ["UR", "1.1", "Enviar Consulta", "1", "1", "1", "1", "1", "1"]
            for i, txt in enumerate(r):
                tk.Label(f, text=txt, bg="white", font=("Times New Roman", 9, "bold")).grid(row=2, column=i, padx=1, pady=1)
                
            tk.Label(f, text="El sistema debe permitir al Visitante enviar una Consulta...", bg="white", font=("Times New Roman", 9), anchor="w").grid(row=3, column=1, columnspan=8, sticky="we", padx=1, pady=1)
            
        elif tipo == "apa":
            tk.Label(win, text="REFERENCIAS BIBLIOGRÁFICAS", font=("Times New Roman", 16, "bold"), bg="white").pack(pady=(50,20), anchor="w", padx=50)
            tk.Label(win, text="American Psychological Association. (2010). Manual de publicaciones de la American...", font=("Arial", 11), bg="white", justify="left").pack(anchor="w", padx=50)

    # ================= MOTOR GRÁFICO A4 =================
    def _construir_motor_grafico(self):
        self.canvas_preview = tk.Canvas(self.tab_preview, bg=self.BG_MAIN, highlightthickness=0)
        scroll_y = ttk.Scrollbar(self.tab_preview, orient="vertical", command=self.canvas_preview.yview)
        self.canvas_preview.configure(yscrollcommand=scroll_y.set)
        scroll_y.pack(side="right", fill="y")
        self.canvas_preview.pack(side="left", fill="both", expand=True)

    def _al_cambiar_pestana(self, event):
        if self.notebook.index("current") == 2:
            grupo = self.ent_grupo.get().strip() or "1"
            webbrowser.open(f"http://localhost:8000/preview-html?doc=proyecto{grupo}")

    def ejecutar_motor_grafico(self):
        self.canvas_preview.delete("all")
        if not self.elementos_documento:
            self.canvas_preview.create_text(300, 100, text="El documento está vacío.", font=('Segoe UI', 14), fill="gray")
            return

        W_PAGINA = 500; H_PAGINA = 707; MARGEN = 45; ESPACIO_ENTRE_PAGINAS = 40
        paginas_totales = 0; y_cursor = MARGEN 
        
        def dibujar_hoja(num_pag):
            y_ini = num_pag * (H_PAGINA + ESPACIO_ENTRE_PAGINAS) + ESPACIO_ENTRE_PAGINAS
            self.canvas_preview.create_rectangle(150+3, y_ini+3, 150+W_PAGINA+3, y_ini+H_PAGINA+3, fill="#D1D5DB", outline="")
            self.canvas_preview.create_rectangle(150, y_ini, 150+W_PAGINA, y_ini+H_PAGINA, fill="white", outline="#D1D5DB")
            self.canvas_preview.create_text(150+W_PAGINA-MARGEN, y_ini+H_PAGINA-20, text=str(num_pag+1), anchor="e", font=('Segoe UI', 8), fill="#9CA3AF")
            return y_ini
        
        y_abs_hoja = dibujar_hoja(paginas_totales)
        cap_num = 0; sub_num = tab_num = fig_num = 0

        for bloque in self.elementos_documento:
            m = bloque["metodo"]; args = bloque["args"]
            altura_estimada = 40
            if m == "texto": altura_estimada = 60
            elif m in ["tabla_simple", "caso_uso", "matriz_ur_esa"]: altura_estimada = 140
            elif m == "figura": altura_estimada = 200
            
            if m == "capitulo" or m == "apa" or (y_cursor + altura_estimada > H_PAGINA - MARGEN):
                if y_cursor > MARGEN: 
                    paginas_totales += 1; y_abs_hoja = dibujar_hoja(paginas_totales); y_cursor = MARGEN
            
            y_real = y_abs_hoja + y_cursor
            x_real = 150 + MARGEN
            w_disp = W_PAGINA - (MARGEN * 2) 
            
            if m == "capitulo":
                cap_num += 1; sub_num = tab_num = fig_num = 0
                self.canvas_preview.create_text(x_real, y_real, text=f"{cap_num}. {args.get('titulo','').upper()}", anchor="nw", font=('Segoe UI', 12, 'bold'), width=w_disp)
                y_cursor += 40
            elif m == "subtitulo":
                sub_num += 1
                self.canvas_preview.create_text(x_real + 15, y_real, text=f"{cap_num}.{sub_num} {args.get('subtitulo','')}", anchor="nw", font=('Segoe UI', 10, 'bold'), width=w_disp-15)
                y_cursor += 30
            elif m == "texto":
                self.canvas_preview.create_text(x_real, y_real, text=f"{args.get('texto','')[:150]}...", anchor="nw", font=('Segoe UI', 9), fill="#4B5563", width=w_disp)
                y_cursor += 60
            elif m == "figura":
                fig_num += 1
                self.canvas_preview.create_rectangle(x_real+30, y_real, x_real+w_disp-30, y_real+120, fill="#DBEAFE", outline="#93C5FD", width=2)
                self.canvas_preview.create_text(x_real+w_disp/2, y_real+60, text=f"🖼️ FIGURA {cap_num}.{fig_num}\n{args.get('leyenda','')}", anchor="center", font=('Segoe UI', 9, 'bold'), fill="#1E40AF")
                y_cursor += 140
            elif m in ["tabla_simple", "caso_uso"]:
                tab_num += 1
                self.canvas_preview.create_text(x_real+w_disp/2, y_real, text=f"TABLA / CASO DE USO {cap_num}.{tab_num}", anchor="center", font=('Segoe UI', 8, 'italic'))
                self.canvas_preview.create_rectangle(x_real, y_real+15, x_real+w_disp, y_real+100, fill="#F9FAFB", outline="#D1D5DB")
                self.canvas_preview.create_rectangle(x_real, y_real+15, x_real+w_disp, y_real+35, fill="#E5E7EB", outline="#D1D5DB") 
                self.canvas_preview.create_text(x_real+w_disp/2, y_real+55, text="📊 DATOS DE LA TABLA", anchor="center", font=('Segoe UI', 10, 'bold'), fill="#6B7280")
                y_cursor += 120
            elif m == "matriz_ur_esa":
                tab_num += 1
                num_cats = len(args.get("categorias", []))
                num_reqs = sum(len(c.get("reqs", [])) for c in args.get("categorias", []))
                self.canvas_preview.create_text(x_real+w_disp/2, y_real, text=f"MATRIZ REQUERIMIENTOS ESA {cap_num}.{tab_num}", anchor="center", font=('Segoe UI', 8, 'italic'))
                self.canvas_preview.create_rectangle(x_real, y_real+15, x_real+w_disp, y_real+100, fill="#F9FAFB", outline="#D1D5DB")
                self.canvas_preview.create_rectangle(x_real, y_real+15, x_real+w_disp, y_real+35, fill="#E5E7EB", outline="#D1D5DB") 
                self.canvas_preview.create_text(x_real+w_disp/2, y_real+55, text=f"📋 MATRIZ UR ({num_cats} Categorías | {num_reqs} Reqs)", anchor="center", font=('Segoe UI', 10, 'bold'), fill="#6B7280")
                y_cursor += 120
            elif m == "apa":
                self.canvas_preview.create_text(x_real, y_real, text="REFERENCIAS BIBLIOGRÁFICAS", anchor="nw", font=('Segoe UI', 12, 'bold'), width=w_disp)
                y_cursor += 40

        altura_total = (paginas_totales + 1) * (H_PAGINA + ESPACIO_ENTRE_PAGINAS) + ESPACIO_ENTRE_PAGINAS
        self.canvas_preview.configure(scrollregion=(0, 0, 800, altura_total))

    # ================= FUNCIONES LÓGICAS (CRUD) =================
    def _mover_arriba(self):
        sel = self.listbox.curselection(); 
        if not sel or sel[0] == 0: return
        idx = sel[0]
        self.elementos_documento[idx], self.elementos_documento[idx-1] = self.elementos_documento[idx-1], self.elementos_documento[idx]
        texto = self.listbox.get(idx)
        self.listbox.delete(idx); self.listbox.insert(idx-1, texto); self.listbox.selection_set(idx-1)

    def _mover_abajo(self):
        sel = self.listbox.curselection()
        if not sel or sel[0] == len(self.elementos_documento)-1: return
        idx = sel[0]
        self.elementos_documento[idx], self.elementos_documento[idx+1] = self.elementos_documento[idx+1], self.elementos_documento[idx]
        texto = self.listbox.get(idx)
        self.listbox.delete(idx); self.listbox.insert(idx+1, texto); self.listbox.selection_set(idx+1)

    def _eliminar_elemento(self):
        sel = self.listbox.curselection()
        if not sel: return
        idx = sel[0]
        self.listbox.delete(idx); self.elementos_documento.pop(idx)

    def _editar_elemento(self):
        sel = self.listbox.curselection()
        if not sel: 
            messagebox.showwarning("Aviso", "Selecciona un elemento de la lista para editarlo.")
            return
        idx = sel[0]
        bloque = self.elementos_documento[idx]
        m = bloque["metodo"]; args = bloque["args"]

        if m == "capitulo": self._dlg_capitulo(idx, args)
        elif m == "subtitulo": self._dlg_subtitulo(idx, args)
        elif m == "texto": self._dlg_texto(idx, args)
        elif m == "definicion": self._dlg_definicion(idx, args)
        elif m == "figura": self._dlg_figura(idx, args)
        elif m == "tabla_simple": self._dlg_tabla_simple(idx, args)
        elif m == "caso_uso": self._dlg_caso_uso(idx, args)
        elif m == "matriz_ur_esa": self._dlg_matriz_ur(idx, args)
        elif m == "apa": self._dlg_apa(idx, args)

    def _abrir_formulario(self, titulo_ventana, campos, callback):
        win = tk.Toplevel(self.root)
        win.title(titulo_ventana); win.geometry("550x450"); win.configure(bg=self.BG_CARD); win.grab_set() 
        tk.Label(win, text=titulo_ventana, font=("Segoe UI", 14, "bold"), bg=self.BG_CARD, fg=self.C_PRIMARY).pack(anchor='w', padx=20, pady=(20,10))
        entradas = []
        for i, (label_txt, tipo, val_defecto) in enumerate(campos):
            tk.Label(win, text=label_txt, font=('Segoe UI', 9, 'bold'), bg=self.BG_CARD, fg=self.C_TEXT).pack(anchor='w', padx=20, pady=(5,2))
            if tipo == "entry":
                ent = ttk.Entry(win, width=65)
                ent.insert(0, str(val_defecto)); ent.pack(padx=20); entradas.append(ent)
            elif tipo == "text":
                txt = tk.Text(win, width=65, height=6, font=('Segoe UI', 9), relief="flat", highlightthickness=1, highlightbackground="#D1D5DB", highlightcolor=self.C_PRIMARY)
                txt.insert("1.0", str(val_defecto)); txt.pack(padx=20); entradas.append(txt)

        def guardar():
            valores = []
            for ent, (_, tipo, _) in zip(entradas, campos):
                if tipo == "entry": valores.append(ent.get())
                elif tipo == "text": valores.append(ent.get("1.0", tk.END).strip())
            callback(valores)
            win.destroy()
        self._btn(win, "💾 Guardar Cambios", guardar, self.C_PRIMARY, self.C_PRIMARY_H, pady=10).pack(pady=20, padx=20, fill='x')

    def _guardar_bloque(self, idx, titulo_lista, tipo_metodo, kwargs):
        if idx is None:
            self.elementos_documento.append({"metodo": tipo_metodo, "args": kwargs})
            self.listbox.insert(tk.END, titulo_lista)
        else:
            self.elementos_documento[idx] = {"metodo": tipo_metodo, "args": kwargs}
            self.listbox.delete(idx); self.listbox.insert(idx, titulo_lista); self.listbox.selection_set(idx)

    # ================= DIÁLOGOS =================
    def _dlg_capitulo(self, idx=None, args=None):
        val = args.get("titulo", "") if args else ""
        self._abrir_formulario("Capítulo", [("Título del Capítulo", "entry", val)], lambda res: self._guardar_bloque(idx, f"📘 CAPÍTULO: {res[0]}", "capitulo", {"titulo": res[0]}))

    def _dlg_subtitulo(self, idx=None, args=None):
        val = args.get("subtitulo", "") if args else ""
        self._abrir_formulario("Subtítulo", [("Subtítulo", "entry", val)], lambda res: self._guardar_bloque(idx, f"   ↳ Subtítulo: {res[0]}", "subtitulo", {"subtitulo": res[0]}))

    def _dlg_texto(self, idx=None, args=None):
        val = args.get("texto", "") if args else ""
        self._abrir_formulario("Texto Libre", [("Escribe tu texto aquí:", "text", val)], lambda res: self._guardar_bloque(idx, f"📝 Texto: {res[0][:40]}...", "texto", {"texto": res[0]}))

    def _dlg_definicion(self, idx=None, args=None):
        v1 = args.get("termino", "") if args else ""; v2 = args.get("descripcion", "") if args else ""
        self._abrir_formulario("Definición / Actor", [("Término", "entry", v1), ("Descripción:", "text", v2)], lambda res: self._guardar_bloque(idx, f"👤 Actor: {res[0]}", "definicion", {"termino": res[0], "descripcion": res[1]}))

    def _dlg_figura(self, idx=None, args=None):
        v1 = args.get("ruta_imagen", "") if args else ""; v2 = args.get("leyenda", "") if args else ""
        self._abrir_formulario("Figura", [("Nombre archivo", "entry", v1), ("Leyenda", "entry", v2)], lambda res: self._guardar_bloque(idx, f"🖼️ Figura: {res[1]}", "figura", {"ruta_imagen": res[0], "leyenda": res[1]}))

    def _dlg_apa(self, idx=None, args=None):
        val = "\n".join(args.get("lista_referencias", [])) if args else ""
        self._abrir_formulario("Bibliografía APA", [("Referencias (Una por línea):", "text", val)], lambda res: self._guardar_bloque(idx, "📚 Bibliografía APA", "apa", {"lista_referencias": res[0].split('\n')}))

    def _dlg_tabla_simple(self, idx=None, args=None):
        win = tk.Toplevel(self.root)
        win.title("Editor de Tabla Dinámica"); win.geometry("800x600"); win.configure(bg=self.BG_CARD); win.grab_set()
        frame_top = tk.Frame(win, bg=self.BG_CARD); frame_top.pack(fill='x', padx=20, pady=10)

        tk.Label(frame_top, text="Leyenda:", bg=self.BG_CARD, font=('Segoe UI', 9, 'bold')).grid(row=0, column=0, sticky='e')
        ent_leyenda = ttk.Entry(frame_top, width=40); ent_leyenda.insert(0, args.get("leyenda", "") if args else ""); ent_leyenda.grid(row=0, column=1, padx=5, pady=5, sticky='w')

        tk.Label(frame_top, text="Color Encabezado:", bg=self.BG_CARD, font=('Segoe UI', 9, 'bold')).grid(row=1, column=0, sticky='e')
        ent_color = ttk.Entry(frame_top, width=15); ent_color.insert(0, args.get("color_encabezado", "D9D9D9") if args else "D9D9D9"); ent_color.grid(row=1, column=1, padx=5, pady=5, sticky='w')

        tk.Label(frame_top, text="Columnas (separadas por coma):", bg=self.BG_CARD, font=('Segoe UI', 9, 'bold')).grid(row=2, column=0, sticky='e')
        ent_headers = ttk.Entry(frame_top, width=60); ent_headers.insert(0, ",".join(args.get("encabezados", ["ID", "Tarea", "Estado"])) if args else "ID, Tarea, Estado"); ent_headers.grid(row=2, column=1, padx=5, pady=5, sticky='w')

        frame_grid_container = tk.Frame(win, bg=self.BG_MAIN, highlightthickness=1, highlightbackground="#D1D5DB"); frame_grid_container.pack(fill='both', expand=True, padx=20, pady=5)
        canvas_grid = tk.Canvas(frame_grid_container, bg=self.BG_MAIN, highlightthickness=0)
        scroll_g = ttk.Scrollbar(frame_grid_container, orient="vertical", command=canvas_grid.yview)
        frame_grid = tk.Frame(canvas_grid, bg=self.BG_MAIN)
        canvas_grid.create_window((0, 0), window=frame_grid, anchor="nw")
        canvas_grid.configure(yscrollcommand=scroll_g.set)
        canvas_grid.pack(side="left", fill="both", expand=True); scroll_g.pack(side="right", fill="y")
        frame_grid.bind("<Configure>", lambda e: canvas_grid.configure(scrollregion=canvas_grid.bbox("all")))

        filas_entries = []

        def renderizar_grid():
            for w in frame_grid.winfo_children(): w.destroy()
            filas_entries.clear()
            headers = [h.strip() for h in ent_headers.get().split(',') if h.strip()]
            for col, h in enumerate(headers): tk.Label(frame_grid, text=h, font=('Segoe UI', 9, 'bold'), bg="#D1D5DB", width=15).grid(row=0, column=col, sticky='nsew', padx=1, pady=1)
            filas_previas = args.get("filas", [["", "", ""]]) if args else [["" for _ in headers]]
            for row_data in filas_previas: add_row(row_data)

        def add_row(datos=None):
            headers = [h.strip() for h in ent_headers.get().split(',') if h.strip()]
            row_idx = len(filas_entries) + 1; current_row_entries = []
            for col in range(len(headers)):
                ent = ttk.Entry(frame_grid, width=15); 
                if datos and col < len(datos): ent.insert(0, datos[col])
                ent.grid(row=row_idx, column=col, padx=1, pady=1, sticky='nsew'); current_row_entries.append(ent)
            filas_entries.append(current_row_entries)

        def actualizar_headers(*_):
            datos_temp = [[e.get() for e in fila] for fila in filas_entries]
            renderizar_grid()
            for r_idx, fila in enumerate(filas_entries):
                for c_idx, ent in enumerate(fila):
                    if r_idx < len(datos_temp) and c_idx < len(datos_temp[r_idx]): ent.insert(0, datos_temp[r_idx][c_idx])

        ent_headers.bind("<FocusOut>", actualizar_headers); ent_headers.bind("<Return>", actualizar_headers)

        frame_btn = tk.Frame(win, bg=self.BG_CARD); frame_btn.pack(fill='x', padx=20, pady=10)
        self._btn(frame_btn, "➕ Añadir Fila", add_row, "#E5E7EB", "#D1D5DB", fg=self.C_TEXT).pack(side=tk.LEFT)

        def guardar():
            leyenda = ent_leyenda.get(); color = ent_color.get(); headers = [h.strip() for h in ent_headers.get().split(',') if h.strip()]
            filas = [[e.get() for e in fila] for fila in filas_entries]
            self._guardar_bloque(idx, f"📊 Tabla: {leyenda}", "tabla_simple", {"leyenda": leyenda, "color_encabezado": color, "encabezados": headers, "filas": filas}); win.destroy()
        self._btn(frame_btn, "💾 Guardar Tabla", guardar, self.C_PRIMARY, self.C_PRIMARY_H).pack(side=tk.RIGHT)
        renderizar_grid()

    def _dlg_caso_uso(self, idx=None, args=None):
        win = tk.Toplevel(self.root)
        win.title("Editor Visual de Caso de Uso")
        win.geometry("750x650")
        win.configure(bg=self.BG_CARD)
        win.grab_set()

        tk.Label(win, text="Formato de Caso de Uso", font=("Segoe UI", 12, "bold"), bg=self.BG_CARD, fg=self.C_PRIMARY).pack(pady=10)

        frame_grid_container = tk.Frame(win, bg=self.BG_MAIN, highlightthickness=1, highlightbackground="#D1D5DB")
        frame_grid_container.pack(fill='both', expand=True, padx=20, pady=10)
        canvas = tk.Canvas(frame_grid_container, bg=self.BG_MAIN, highlightthickness=0)
        scroll = ttk.Scrollbar(frame_grid_container, orient="vertical", command=canvas.yview)
        frame_grid = tk.Frame(canvas, bg=self.BG_MAIN)
        canvas.create_window((0, 0), window=frame_grid, anchor="nw")
        canvas.configure(yscrollcommand=scroll.set)
        canvas.pack(side="left", fill="both", expand=True); scroll.pack(side="right", fill="y")
        frame_grid.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

        def crear_fila(row, label, height=1, val=""):
            tk.Label(frame_grid, text=label, font=('Segoe UI', 9, 'bold'), bg="#F3F4F6", anchor='w', width=20, padx=10).grid(row=row, column=0, sticky='nsew', padx=1, pady=1)
            if height == 1:
                ent = ttk.Entry(frame_grid, width=60); ent.insert(0, val); ent.grid(row=row, column=1, sticky='nsew', padx=1, pady=1, ipady=4); return ent
            else:
                txt = tk.Text(frame_grid, height=height, width=60, font=('Segoe UI', 9), relief="flat", padx=5, pady=5); txt.insert("1.0", val); txt.grid(row=row, column=1, sticky='nsew', padx=1, pady=1); return txt

        ent_nombre = crear_fila(0, "Nombre", 1, args.get("nombre", "") if args else "")
        ent_actores = crear_fila(1, "Actor(es)", 1, args.get("actores", "") if args else "")
        txt_resumen = crear_fila(2, "Resumen", 2, args.get("resumen", "") if args else "")
        ent_frec = crear_fila(3, "Frecuencia", 1, args.get("frecuencia", "") if args else "")
        txt_pre = crear_fila(4, "Precondiciones", 2, args.get("precondiciones", "") if args else "")
        txt_desc = crear_fila(5, "Descripción", 4, args.get("descripcion", "") if args else "")
        txt_exc = crear_fila(6, "Excepciones", 2, args.get("excepciones", "") if args else "")
        txt_pos = crear_fila(7, "Poscondiciones", 2, args.get("poscondiciones", "") if args else "")
        txt_dep = crear_fila(8, "Dependencias", 2, args.get("dependencias", "") if args else "")

        frame_grid.columnconfigure(1, weight=1)

        def guardar():
            self._guardar_bloque(idx, f"⚙️ Caso de Uso: {ent_nombre.get()}", "caso_uso", {
                "nombre": ent_nombre.get(), "actores": ent_actores.get(), "resumen": txt_resumen.get("1.0", tk.END).strip(),
                "frecuencia": ent_frec.get(), "precondiciones": txt_pre.get("1.0", tk.END).strip(),
                "descripcion": txt_desc.get("1.0", tk.END).strip(), "excepciones": txt_exc.get("1.0", tk.END).strip(),
                "poscondiciones": txt_pos.get("1.0", tk.END).strip(), "dependencias": txt_dep.get("1.0", tk.END).strip()
            }); win.destroy()

        self._btn(win, "💾 Guardar Caso de Uso", guardar, self.C_PRIMARY, self.C_PRIMARY_H, font=('Segoe UI', 10, 'bold')).pack(pady=15, ipadx=10)

    def _dlg_matriz_ur(self, idx=None, args=None):
        win = tk.Toplevel(self.root)
        win.title("Editor de Matriz UR (Estándar ESA)")
        win.geometry("1000x700")
        win.configure(bg=self.BG_CARD)
        win.grab_set()

        tk.Label(win, text="Matriz de Requerimientos UR Continua", font=("Segoe UI", 14, "bold"), bg=self.BG_CARD, fg=self.C_PRIMARY).pack(pady=10)

        frame_grid_container = tk.Frame(win, bg=self.BG_MAIN, highlightthickness=1, highlightbackground="#D1D5DB")
        frame_grid_container.pack(fill='both', expand=True, padx=20, pady=5)

        canvas = tk.Canvas(frame_grid_container, bg=self.BG_MAIN, highlightthickness=0)
        scroll = ttk.Scrollbar(frame_grid_container, orient="vertical", command=canvas.yview)
        frame_canvas = tk.Frame(canvas, bg=self.BG_MAIN)

        canvas.create_window((0, 0), window=frame_canvas, anchor="nw")
        canvas.configure(yscrollcommand=scroll.set)
        canvas.pack(side="left", fill="both", expand=True)
        scroll.pack(side="right", fill="y")
        frame_canvas.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

        frame_head = tk.Frame(frame_canvas, bg=self.BG_CARD)
        frame_head.pack(fill='x', pady=(0, 10))
        headers = ["Tipo", "ID UR", "Nombre del Requisito", "Nec", "Prio", "Est", "Cla", "Ver", "Fue", "Acciones"]
        widths = [5, 6, 40, 4, 4, 4, 4, 4, 4, 15]
        for col, (h, w) in enumerate(zip(headers, widths)):
            tk.Label(frame_head, text=h, font=('Segoe UI', 8, 'bold'), bg="#D1D5DB", width=w).grid(row=0, column=col, sticky='nsew', padx=1)

        cats_widgets = []

        def add_category(cat_data=None):
            if cat_data is None: cat_data = {}
            f_cat = tk.Frame(frame_canvas, bg="#F3F4F6", highlightthickness=1, highlightbackground="#9CA3AF", pady=5, padx=5)
            f_cat.pack(fill='x', pady=10)

            f_cat_head = tk.Frame(f_cat, bg="#F3F4F6")
            f_cat_head.pack(fill='x')
            tk.Label(f_cat_head, text=f"Categoría {len(cats_widgets)+1}:", font=('Segoe UI', 10, 'bold'), bg="#F3F4F6", fg=self.C_PRIMARY).pack(side=tk.LEFT)
            e_cat_name = ttk.Entry(f_cat_head, width=60)
            e_cat_name.insert(0, cat_data.get("nombre", ""))
            e_cat_name.pack(side=tk.LEFT, padx=10)

            cat_dict = {"frame": f_cat, "e_nombre": e_cat_name, "reqs": []}
            btn_del_cat = tk.Button(f_cat_head, text="✖ Eliminar Categoría", fg="white", bg=self.C_DANGER, relief="flat", cursor="hand2", command=lambda: del_cat(cat_dict))
            btn_del_cat.pack(side=tk.RIGHT)

            f_reqs_container = tk.Frame(f_cat, bg="#F3F4F6")
            f_reqs_container.pack(fill='x', pady=5)

            def add_req(req_data=None):
                if req_data is None: req_data = {}
                f_req = tk.Frame(f_reqs_container, bg=self.BG_CARD, highlightthickness=1, highlightbackground="#D1D5DB", pady=5, padx=5)
                f_req.pack(fill='x', pady=2)

                f_r1 = tk.Frame(f_req, bg=self.BG_CARD)
                f_r1.pack(fill='x')

                e_tipo = ttk.Entry(f_r1, width=6, justify='center'); e_tipo.insert(0, req_data.get("tipo", "UR")); e_tipo.pack(side=tk.LEFT, padx=1)
                e_id = ttk.Entry(f_r1, width=7, justify='center'); e_id.insert(0, req_data.get("id", "")); e_id.pack(side=tk.LEFT, padx=1)
                e_nom = ttk.Entry(f_r1, width=50); e_nom.insert(0, req_data.get("nombre", "")); e_nom.pack(side=tk.LEFT, padx=1)

                e_pts = []
                pts_data = req_data.get("puntos", ["", "", "", "", "", ""])
                while len(pts_data) < 6: pts_data.append("")
                for i in range(6):
                    ep = ttk.Entry(f_r1, width=5, justify='center')
                    ep.insert(0, pts_data[i])
                    ep.pack(side=tk.LEFT, padx=1)
                    e_pts.append(ep)

                req_dict_item = {"frame": f_req, "tipo": e_tipo, "id": e_id, "nombre": e_nom, "pts": e_pts}

                btn_del_req = tk.Button(f_r1, text="✖", fg="white", bg=self.C_DANGER, relief="flat", command=lambda: del_req(cat_dict, req_dict_item))
                btn_del_req.pack(side=tk.RIGHT, padx=5)

                f_r2 = tk.Frame(f_req, bg=self.BG_CARD)
                f_r2.pack(fill='x', pady=(5,0))
                tk.Label(f_r2, text="Desc:", bg=self.BG_CARD, font=('Segoe UI', 8, 'bold')).pack(side=tk.LEFT)
                t_desc = tk.Text(f_r2, height=2, font=('Segoe UI', 9), relief="flat", highlightthickness=1, highlightbackground="#E5E7EB")
                t_desc.insert("1.0", req_data.get("descripcion", ""))
                t_desc.pack(fill='x', expand=True, padx=5)
                req_dict_item["desc"] = t_desc

                cat_dict["reqs"].append(req_dict_item)

            def del_req(cat_d, req_d):
                req_d["frame"].destroy(); cat_d["reqs"].remove(req_d)

            if "reqs" in cat_data:
                for r in cat_data["reqs"]: add_req(r)
            else:
                add_req()

            self._btn(f_cat, "➕ Añadir Requisito a esta Categoría", add_req, "#E5E7EB", "#D1D5DB", fg=self.C_TEXT, pady=5).pack(pady=5)
            cats_widgets.append(cat_dict)

        def del_cat(cat_d):
            cat_d["frame"].destroy(); cats_widgets.remove(cat_d)

        if args and "categorias" in args:
            for c in args["categorias"]: add_category(c)
        else:
            add_category({"nombre": "Formulario de Contacto: Visitante", "reqs": [{"tipo":"UR", "id":"1.1", "nombre":"Enviar Consulta", "puntos":["1","1","1","1","1","1"], "descripcion":"El sistema debe permitir..."}]})

        frame_btn = tk.Frame(win, bg=self.BG_CARD)
        frame_btn.pack(fill='x', padx=20, pady=10)
        self._btn(frame_btn, "➕ Añadir Nueva Categoría UR", add_category, self.C_SUCCESS, self.C_SUCCESS_H).pack(side=tk.LEFT)

        def guardar():
            datos_finales = []
            for cw in cats_widgets:
                nombre_cat = cw["e_nombre"].get().strip()
                reqs_lista = []
                for rw in cw["reqs"]:
                    reqs_lista.append({
                        "tipo": rw["tipo"].get().strip(), "id": rw["id"].get().strip(),
                        "nombre": rw["nombre"].get().strip(), "puntos": [e.get().strip() for e in rw["pts"]],
                        "descripcion": rw["desc"].get("1.0", tk.END).strip()
                    })
                datos_finales.append({"nombre": nombre_cat, "reqs": reqs_lista})

            self._guardar_bloque(idx, f"📋 MATRIZ UR ESA ({len(datos_finales)} Categorías)", "matriz_ur_esa", {"categorias": datos_finales})
            win.destroy()

        self._btn(frame_btn, "💾 Guardar Matriz Total", guardar, self.C_PRIMARY, self.C_PRIMARY_H).pack(side=tk.RIGHT)

    # ================== COMPILADOR ==================
    def generar_documento(self):
        if not self.elementos_documento:
            messagebox.showwarning("Aviso", "Añade bloques desde el panel izquierdo primero.")
            return

        titulo = self.ent_titulo.get(); grupo = self.ent_grupo.get()
        autores = [a for a in self.txt_autores.get('1.0', tk.END).strip().split('\n') if a.strip()]

        try:
            inf = GeneradorInforme(titulo, grupo)
            inf.generar_portada(self.ent_logo.get(), self.ent_facultad.get(), self.ent_curso.get(), autores, self.ent_profesor.get(), self.ent_ciudad.get(), self.ent_anio.get())
            inf.estructurar_indices()
            
            cap_num = 0; sub_num = tab_num = fig_num = 0

            for bloque in self.elementos_documento:
                m = bloque["metodo"]; args = dict(bloque["args"])
                
                if m == "capitulo": 
                    cap_num += 1; sub_num = tab_num = fig_num = 0
                    args["titulo"] = f"{cap_num}. {args['titulo']}"
                    inf.agregar_capitulo(**args)
                elif m == "subtitulo": 
                    sub_num += 1; args["subtitulo"] = f"{cap_num}.{sub_num} {args['subtitulo']}"
                    inf.agregar_subtitulo(**args)
                elif m == "figura": 
                    fig_num += 1; args["correlativo"] = f"{cap_num}.{fig_num}"
                    inf.agregar_figura(**args)
                elif m == "tabla_simple": 
                    tab_num += 1; args["correlativo"] = f"{cap_num}.{tab_num}"
                    inf.agregar_tabla(**args)
                elif m == "caso_uso": 
                    tab_num += 1; args["correlativo"] = f"{cap_num}.{tab_num}"
                    inf.agregar_caso_uso_extendido(**args)
                elif m == "matriz_ur_esa": 
                    tab_num += 1; args["correlativo"] = f"{cap_num}.{tab_num}"
                    inf.agregar_matriz_ur_esa(**args)
                elif m == "texto": inf.agregar_texto_libre(**args)
                elif m == "definicion": inf.agregar_definicion(**args)
                elif m == "apa": inf.agregar_bibliografia_apa(**args)

            nombre_archivo = inf.compilar()
            messagebox.showinfo("¡Éxito!", f"Documento '{nombre_archivo}' generado.\n\nRecuerda apretar Ctrl+E y F9 en Word.")
        
        except Exception as e:
            messagebox.showerror("Error", f"Error al generar:\n{str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = AppGeneradorGUI(root)
    root.mainloop()